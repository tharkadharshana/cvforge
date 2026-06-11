import io
from conftest import auth_headers
from app.schemas import CVData, Contact, Experience
from app.cv import render
from app.cv.extract import extract_text, UnsupportedFile
import pytest


def test_audit_trail_and_admin_access(client):
    H = auth_headers(client, email="trace@test.com")
    client.post("/auth/login", data={"username": "trace@test.com", "password": "wrong"})  # failed login
    admin = auth_headers(client, email="admin@test.com")

    # non-admin blocked
    assert client.get("/admin/users", headers=H).status_code == 403

    # admin investigation view
    users = client.get("/admin/users?query=trace", headers=admin).json()
    uid = users[0]["id"]
    detail = client.get(f"/admin/users/{uid}", headers=admin).json()
    events = {e["event"] for e in detail["audit"]}
    assert "register" in events
    assert "login" in events

    # failed login is captured
    fails = client.get("/admin/audit?event=login_failed", headers=admin).json()
    assert len(fails) >= 1


def test_admin_credit_adjust_is_audited(client):
    auth_headers(client, email="adj@test.com")
    admin = auth_headers(client, email="admin@test.com")
    uid = client.get("/admin/users?query=adj", headers=admin).json()[0]["id"]
    out = client.post(f"/admin/users/{uid}/credits", headers=admin, json={"delta": 25, "reason": "goodwill"}).json()
    assert out["credits"] == 2 + 25
    adj = client.get(f"/admin/audit?user_id={uid}&event=admin_credit_adjust", headers=admin).json()
    assert len(adj) == 1 and adj[0]["meta"]["delta"] == 25


def _sample_cv():
    return CVData(
        contact=Contact(full_name="Jane Doe", email="jane@test.com", location="Colombo"),
        summary="Engineer.",
        skills={"Languages": ["Python", "Go"]},
        experience=[Experience(title="Eng", company="Acme", start="2020", end="2024", bullets=["Built things"])],
        languages=["English"],
    )


def test_renderers_produce_valid_bytes():
    cv = _sample_cv()
    docx = render.render_docx(cv)
    pdf = render.render_pdf(cv)
    cover_pdf = render.render_cover_letter_pdf("Para one.\n\nPara two.", "a@b.com", "Jane")
    cover_docx = render.render_cover_letter_docx("Para one.\n\nPara two.", "Jane")
    assert docx[:2] == b"PK"          # docx is a zip
    assert pdf[:4] == b"%PDF"
    assert cover_pdf[:4] == b"%PDF"
    assert cover_docx[:2] == b"PK"


def test_text_extraction():
    txt = extract_text("cv.txt", b"John Smith\nData engineer with strong Python and SQL skills and delivery record.")
    assert "Data engineer" in txt
    # docx
    from docx import Document
    d = Document()
    d.add_paragraph("Jane Doe")
    d.add_paragraph("Senior Backend Engineer with FastAPI and PostgreSQL experience building systems.")
    buf = io.BytesIO(); d.save(buf)
    assert "Backend Engineer" in extract_text("cv.docx", buf.getvalue())
    # legacy .doc rejected
    with pytest.raises(UnsupportedFile):
        extract_text("cv.doc", b"x" * 50)
