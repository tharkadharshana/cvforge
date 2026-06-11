from __future__ import annotations
import io
from ..logging_config import get_logger

log = get_logger("extract")

MAX_BYTES = 10 * 1024 * 1024  # 10 MB


class UnsupportedFile(Exception):
    pass


def extract_text(filename: str, content: bytes, content_type: str = "") -> str:
    name = (filename or "").lower()
    size = len(content)
    log.info("extract: file=%r type=%r bytes=%d", filename, content_type, size)
    if size == 0:
        raise UnsupportedFile("empty file")
    if size > MAX_BYTES:
        raise UnsupportedFile(f"file too large ({size} bytes, max {MAX_BYTES})")

    if name.endswith(".txt") or content_type.startswith("text/"):
        text = _from_txt(content)
    elif name.endswith(".pdf") or content_type == "application/pdf":
        text = _from_pdf(content)
    elif name.endswith(".docx") or content_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        text = _from_docx(content)
    elif name.endswith(".doc"):
        raise UnsupportedFile("legacy .doc not supported — save as .docx or PDF and re-upload")
    else:
        raise UnsupportedFile(f"unsupported file type: {filename or content_type}")

    text = text.strip()
    log.info("extract: ok %d chars", len(text))
    if len(text) < 20:
        raise UnsupportedFile("could not extract readable text (scanned image PDF? try a text PDF or paste text)")
    return text


def _from_txt(content: bytes) -> str:
    for enc in ("utf-8", "latin-1"):
        try:
            return content.decode(enc)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", "replace")


def _from_pdf(content: bytes) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        log.info("extract: pdf pages=%d", len(pdf.pages))
        for i, page in enumerate(pdf.pages):
            t = page.extract_text() or ""
            parts.append(t)
            log.debug("extract: pdf page %d -> %d chars", i + 1, len(t))
    return "\n".join(parts)


def _from_docx(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # include table cell text (CVs often use tables)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("  ".join(cells))
    log.info("extract: docx paragraphs=%d tables=%d", len(doc.paragraphs), len(doc.tables))
    return "\n".join(parts)
