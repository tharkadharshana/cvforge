from __future__ import annotations
import io
from ..schemas import CVData

# docx/fpdf are only needed for download endpoints, not the rest of the API -
# import them lazily inside each render function so other routes (e.g.
# billing) don't pay for them on a cold start.

# ATS rules followed everywhere: single column, standard headings, real text,
# no tables/text-boxes/graphics, common fonts, simple bullets.

SECTION_ORDER_HEADINGS = {
    "summary": "PROFESSIONAL SUMMARY",
    "skills": "SKILLS",
    "experience": "EXPERIENCE",
    "projects": "PROJECTS",
    "education": "EDUCATION",
    "certifications": "CERTIFICATIONS",
    "awards": "AWARDS",
    "languages": "LANGUAGES",
}


def _contact_line(cv: CVData) -> str:
    c = cv.contact
    parts = [c.email, c.phone, c.location, c.linkedin, c.github, c.website]
    return "  |  ".join([p for p in parts if p])


# Default ATS-safe style. render_pdf/render_docx accept an optional `style` dict
# (from cv/templates.resolve_style) that overrides these; anything missing falls
# back here, so callers that pass nothing get the original, byte-stable output.
_DEFAULT_STYLE = {
    "font": "Helvetica", "accent": "#1A1A1A", "heading": "rule",
    "name_size": 18, "body_size": 10.5, "columns": 1,
}

# fpdf core fonts / docx safe fonts. Non-core template fonts (e.g. Inter) map to a
# core font here so ats_safe rendering never fails on a missing font.
_PDF_FONT = {"Helvetica": "Helvetica", "Times": "Times", "Courier": "Courier"}
_DOCX_FONT = {"Helvetica": "Calibri", "Times": "Cambria", "Courier": "Consolas"}


def _hex_rgb(h: str) -> tuple[int, int, int]:
    h = (h or "#1A1A1A").lstrip("#")
    if len(h) != 6:
        h = "1A1A1A"
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _style(style: dict | None) -> dict:
    s = dict(_DEFAULT_STYLE)
    if style:
        s.update({k: v for k, v in style.items() if v is not None})
    return s


# ---------------- DOCX ----------------
def render_docx(cv: CVData, style: dict | None = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    s = _style(style)
    accent = RGBColor(*_hex_rgb(s["accent"]))
    body_font = _DOCX_FONT.get(s["font"], "Calibri")
    caps = s["heading"] == "caps"

    doc = Document()
    base = doc.styles["Normal"]
    base.font.name = body_font
    base.font.size = Pt(float(s["body_size"]))

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(cv.contact.full_name or "")
    run.bold = True
    run.font.size = Pt(float(s["name_size"]))
    run.font.color.rgb = accent

    contact = doc.add_paragraph(_contact_line(cv))
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in contact.runs:
        r.font.size = Pt(9.5)

    def heading(text: str):
        p = doc.add_paragraph()
        run = p.add_run(text.upper() if caps else text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = accent
        p.space_before = Pt(8)

    def bullet(text: str):
        doc.add_paragraph(text, style="List Bullet")

    if cv.summary:
        heading(SECTION_ORDER_HEADINGS["summary"])
        doc.add_paragraph(cv.summary)

    if cv.skills:
        heading(SECTION_ORDER_HEADINGS["skills"])
        for cat, items in cv.skills.items():
            p = doc.add_paragraph()
            p.add_run(f"{cat}: ").bold = True
            p.add_run(", ".join(items))

    if cv.experience:
        heading(SECTION_ORDER_HEADINGS["experience"])
        for e in cv.experience:
            p = doc.add_paragraph()
            p.add_run(f"{e.title}, {e.company}").bold = True
            meta = "  ".join([x for x in [e.location, f"{e.start} - {e.end}".strip(" -")] if x])
            if meta:
                p.add_run(f"   {meta}").italic = True
            for b in e.bullets:
                bullet(b)

    if cv.projects:
        heading(SECTION_ORDER_HEADINGS["projects"])
        for pr in cv.projects:
            p = doc.add_paragraph()
            p.add_run(pr.name).bold = True
            if pr.tech:
                p.add_run(f"  ({', '.join(pr.tech)})").italic = True
            if pr.description:
                doc.add_paragraph(pr.description)
            for b in pr.bullets:
                bullet(b)

    if cv.education:
        heading(SECTION_ORDER_HEADINGS["education"])
        for ed in cv.education:
            p = doc.add_paragraph()
            p.add_run(f"{ed.degree}, {ed.institution}").bold = True
            meta = "  ".join([x for x in [ed.location, f"{ed.start} - {ed.end}".strip(" -")] if x])
            if meta:
                p.add_run(f"   {meta}").italic = True
            for d in ed.details:
                bullet(d)

    if cv.certifications:
        heading(SECTION_ORDER_HEADINGS["certifications"])
        for c in cv.certifications:
            bullet(c)

    if cv.awards:
        heading(SECTION_ORDER_HEADINGS["awards"])
        for a in cv.awards:
            bullet(a)

    if cv.languages:
        heading(SECTION_ORDER_HEADINGS["languages"])
        doc.add_paragraph(", ".join(cv.languages))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------- PDF (pure python, single column) ----------------
def _txt(s: str) -> str:
    # fpdf2 core fonts are latin-1; replace common unicode to keep it robust
    return (s.replace("\u2013", "-").replace("\u2014", "-")
             .replace("\u2018", "'").replace("\u2019", "'")
             .replace("\u201c", '"').replace("\u201d", '"')
             .encode("latin-1", "replace").decode("latin-1"))


def render_pdf(cv: CVData, style: dict | None = None) -> bytes:
    from fpdf import FPDF

    s = _style(style)
    font = _PDF_FONT.get(s["font"], "Helvetica")
    accent = _hex_rgb(s["accent"])
    caps = s["heading"] == "caps"
    body_size = float(s["body_size"])

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(18, 15, 18)

    pdf.set_text_color(*accent)
    pdf.set_font(font, "B", float(s["name_size"]))
    pdf.cell(0, 9, _txt(cv.contact.full_name or ""), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    pdf.set_font(font, "", 9)
    pdf.cell(0, 5, _txt(_contact_line(cv)), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    def heading(text: str):
        pdf.ln(2)
        pdf.set_x(pdf.l_margin)
        pdf.set_text_color(*accent)
        pdf.set_font(font, "B", 12)
        pdf.cell(0, 6, _txt(text.upper() if caps else text), new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        if s["heading"] != "plain":
            pdf.set_draw_color(*accent)
            y = pdf.get_y()
            pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
        pdf.ln(1)

    def body(text: str, bold=False, italic=False, size=None):
        pdf.set_x(pdf.l_margin)
        pdf.set_font(font, ("B" if bold else "") + ("I" if italic else ""), size or body_size)
        pdf.multi_cell(0, 5, _txt(text))

    def bullet(text: str):
        pdf.set_x(pdf.l_margin)
        pdf.set_font(font, "", body_size)
        pdf.multi_cell(0, 5, _txt(f"- {text}"))

    if cv.summary:
        heading("PROFESSIONAL SUMMARY")
        body(cv.summary)

    if cv.skills:
        heading("SKILLS")
        for cat, items in cv.skills.items():
            body(f"{cat}: {', '.join(items)}")

    if cv.experience:
        heading("EXPERIENCE")
        for e in cv.experience:
            body(f"{e.title}, {e.company}", bold=True)
            meta = "  ".join([x for x in [e.location, f"{e.start} - {e.end}".strip(" -")] if x])
            if meta:
                body(meta, italic=True, size=9)
            for b in e.bullets:
                bullet(b)
            pdf.ln(1)

    if cv.projects:
        heading("PROJECTS")
        for pr in cv.projects:
            tech = f" ({', '.join(pr.tech)})" if pr.tech else ""
            body(f"{pr.name}{tech}", bold=True)
            if pr.description:
                body(pr.description)
            for b in pr.bullets:
                bullet(b)
            pdf.ln(1)

    if cv.education:
        heading("EDUCATION")
        for ed in cv.education:
            body(f"{ed.degree}, {ed.institution}", bold=True)
            meta = "  ".join([x for x in [ed.location, f"{ed.start} - {ed.end}".strip(" -")] if x])
            if meta:
                body(meta, italic=True, size=9)
            for d in ed.details:
                bullet(d)

    if cv.certifications:
        heading("CERTIFICATIONS")
        for c in cv.certifications:
            bullet(c)

    if cv.awards:
        heading("AWARDS")
        for a in cv.awards:
            bullet(a)

    if cv.languages:
        heading("LANGUAGES")
        body(", ".join(cv.languages))

    out = pdf.output()
    return bytes(out)


def render_cover_letter_pdf(text: str, contact_line: str = "", name: str = "") -> bytes:
    from fpdf import FPDF

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(20, 18, 20)
    if name:
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 7, _txt(name), new_x="LMARGIN", new_y="NEXT")
    if contact_line:
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(0, 5, _txt(contact_line), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font("Helvetica", "", 11)
    for para in text.split("\n\n"):
        if para.strip():
            pdf.multi_cell(0, 6, _txt(para.strip()))
            pdf.ln(2)
    return bytes(pdf.output())


def render_cover_letter_docx(text: str, name: str = "") -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    if name:
        p = doc.add_paragraph()
        p.add_run(name).bold = True
    for para in text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
